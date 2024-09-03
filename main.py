from docx import Document
from docx.shared import Pt
import random
from faker import Faker

# Initialize Faker for generating random data
fake = Faker()


# Function to create a random table
def create_random_table(document):
    table = document.add_table(rows=9, cols=7)
    table.style = 'Table Grid'

    # Add header row
    headers = ["No.", "Name", "Age", "Department", "DOA", "Fee", "Gender"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Add data rows with random values
    departments = ["Computer", "History", "Hindi"]
    genders = ["Male", "Female"]

    for i in range(1, 9):
        table.cell(i, 0).text = str(i)
        table.cell(i, 1).text = fake.first_name()
        table.cell(i, 2).text = str(random.randint(20, 35))
        table.cell(i, 3).text = random.choice(departments)
        table.cell(i, 4).text = fake.date_of_birth(tzinfo=None, minimum_age=20, maximum_age=30).strftime("%d/%m/%y")
        table.cell(i, 5).text = str(random.randint(100, 400))
        table.cell(i, 6).text = random.choice(genders)


# Function to create a Word document with a random table and lorem ipsum text
def create_word_document(file_name):
    document = Document()

    # Add a title
    document.add_heading('Random Data Table', 0)

    # Add a random paragraph of Lorem Ipsum text
    document.add_paragraph(fake.text(max_nb_chars=200))

    # Add the table with random data
    create_random_table(document)

    # Save the document
    document.save(file_name)


# Create 10 Word documents with random data
for i in range(1, 11):
    file_name = f'random_table_{i}.docx'
    create_word_document(file_name)

print("10 Word documents created with random tables and Lorem Ipsum text.")
